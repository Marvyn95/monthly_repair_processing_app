import streamlit as st
st.set_page_config(page_title="Fleet Repair Monthly", layout="wide")  # <--- add this as the first Streamlit call
import pandas as pd
import datetime
import openpyxl
import os
from openpyxl import load_workbook
from openpyxl.styles import Alignment
from openpyxl.styles import Border, Side
from docx import Document
from docx.shared import Pt, RGBColor  # add RGBColor to your import
from docx.enum.text import WD_ALIGN_PARAGRAPH
from io import BytesIO
from num2words import num2words

st.markdown(
    """
    <style>
    html, body, [class*="css"]  {
        font-size: 14px !important;
    }
    .stDataFrame, .stTable, .stMarkdown, .stTextInput, .stSelectbox, .stButton, .stForm, .stDateInput, .stNumberInput {
        font-size: 14px !important;
    }
    </style>
    """,
    unsafe_allow_html=True
)

sale_entry, sales = st.columns([5, 8])

with sale_entry:
    form_values = {}
    with st.form("sale_form", clear_on_submit=True):
        areas_df = pd.read_excel("static/areas.xlsx")
        area_list = areas_df['Area'].dropna().unique().tolist()

        vehicle_df = pd.read_excel("static/vehicles.xlsx")
        vehicle_list = vehicle_df['Vehicle'].dropna().unique().tolist()

        col1, col2, col3 = st.columns([3,4,3])

        with col1:
            if area_list:
                form_values["area"] = st.selectbox("Area", area_list, key="area_select")
            else:
                form_values["area"] = st.text_input("Area", key="area_text")

        with col2:
            if vehicle_list:
                form_values["vehicle"] = st.selectbox("Vehicle", vehicle_list, key="vehicle_select")
            else:
                form_values["vehicle"] = st.text_input("Vehicle", key="vehicle_id")

        with col3:
            first_day_prev_month = (datetime.date.today().replace(day=1) - datetime.timedelta(days=1)).replace(day=1)
            form_values["date_of_repair"] = st.date_input("Date of Repair", key="date_of_repair", value=first_day_prev_month)

        if "repair_rows" not in st.session_state:
            st.session_state.repair_rows = 5

        for i in range(st.session_state.repair_rows):
            col1, col2 = st.columns(2)

            with col1:
                form_values[f"repair_description_{i}"] = st.text_input(f"Repair Description {i+1}", key=f"repair_description_{i}")

            with col2:
                form_values[f"cost_{i}"] = st.number_input(f"Cost {i+1}", min_value=0, key=f"cost_{i}", value=None, placeholder="")

        submit_pressed = st.form_submit_button("Submit Repair Entry")
        
        if submit_pressed:

            repairs_file_path = os.path.join("static", "repairs_excel.xlsx")
            os.makedirs(os.path.dirname(repairs_file_path), exist_ok=True)

            if not os.path.exists(repairs_file_path):
                empty_df = pd.DataFrame(columns=["No.", "Area", "Vehicle ID", "Date", "Description", "Cost (ugx)"])
                today_str = datetime.date.today().strftime("%d-%m-%Y")
                sheet_name=f"{today_str}"
                with pd.ExcelWriter(repairs_file_path, engine="openpyxl") as writer:
                    empty_df.to_excel(writer, index=False, sheet_name=sheet_name)

            today_str = datetime.date.today().strftime("%d-%m-%Y")
            sheet_name=f"{today_str}"

            repairs_excel_df = pd.read_excel(repairs_file_path, sheet_name=sheet_name)
            num_vehicle_entries = repairs_excel_df["Vehicle ID"].nunique()

            rows = []
            for i in range(st.session_state.repair_rows):
                if i == 0:
                    new_row = {
                        "No.": "",
                        "Area": form_values.get("area"),
                        "Vehicle ID": form_values.get("vehicle"),
                        "Date": pd.to_datetime(form_values.get("date_of_repair")).strftime("%d-%b-%Y")
                    }
                    desc = (form_values.get(f'repair_description_{i}') or '').strip()
                    cost = form_values.get(f'cost_{i}')
                    if desc not in [None, ""] or cost not in [None, 0]:
                        new_row["Description"] = desc
                        new_row["Cost (ugx)"] = f'{cost:,}'
                        rows.append(new_row)
                elif i > 0:
                    desc = (form_values.get(f'repair_description_{i}') or '').strip()
                    cost = form_values.get(f'cost_{i}')
                    if desc not in [None, ""] or cost not in [None, 0]:
                        new_row = {
                            "No.": "",
                            "Area": "",
                            "Vehicle ID": "",
                            "Date": "",
                            "Description": desc,
                            "Cost (ugx)": f'{cost:,}'
                        }
                        rows.append(new_row)
            total_cost = sum([(form_values.get(f'cost_{j}') or 0) for j in range(st.session_state.repair_rows)])
            rows.append({
                "No.": "",
                "Area": "",
                "Vehicle ID": "",
                "Date": "",
                "Description": "Total Cost (ugx)",
                "Cost (ugx)": f'{total_cost:,}'
            })
            rows_df = pd.DataFrame(rows)
            monthly_repairs_df = pd.concat([repairs_excel_df, rows_df], ignore_index=True)

            count = 0
            for i in monthly_repairs_df.index:
                if pd.notna(monthly_repairs_df.at[i, "Vehicle ID"]) and str(monthly_repairs_df.at[i, "Vehicle ID"]).strip() != "":
                    monthly_repairs_df.at[i, "No."] = count + 1
                    count += 1
                else:
                    monthly_repairs_df.at[i, "No."] = ""

            with pd.ExcelWriter(repairs_file_path, engine="openpyxl", mode="a", if_sheet_exists="replace") as writer:
                monthly_repairs_df.to_excel(writer, index=False, sheet_name=sheet_name)

            work_book = load_workbook(repairs_file_path)
            ws = work_book[sheet_name]

            thin_border = Border(
                left=Side(style='thin'),
                right=Side(style='thin'),
                top=Side(style='thin'),
                bottom=Side(style='thin')
            )
            col_max_width = {}

            for row in ws.iter_rows():
                for cell in row:
                    cell.alignment = Alignment(wrap_text=True, vertical="top")
                    cell.border = thin_border
                    val = str(cell.value) if cell.value is not None else ""
                    col_letter = cell.column_letter
                    col_max_width[col_letter] = max(col_max_width.get(col_letter, 0), len(val))

            # Set column widths based on max content length
            for col_letter, max_len in col_max_width.items():
                ws.column_dimensions[col_letter].width = max(15, min(max_len + 3, 60))  # reasonable min/max

            work_book.save(repairs_file_path)

            st.success("Repair entry submitted!")
            st.session_state.clear()
            st.session_state.repair_rows = 5

    col1, col2, col3, col4, col5 = st.columns([1, 1, 1, 1, 1])
    
    with col1:
        def add_row():
            st.session_state.repair_rows += 1
        st.button("Add", key=f"add_{i}", on_click=add_row)
    
    with col2:
        def remove_row():
            if st.session_state.repair_rows > 1:
                st.session_state.repair_rows -= 1
        st.button("Remove", key=f"remove_{i}", on_click=remove_row)


with sales:
    st.markdown("#### Sales Recorded")

    file_path = os.path.join("static", "repairs_excel.xlsx")
    os.makedirs(os.path.dirname(file_path), exist_ok=True)

    if not os.path.exists(file_path):
        empty_df = pd.DataFrame(columns=["No.", "Area", "Vehicle ID", "Date", "Description", "Cost (ugx)"])
        today_str = datetime.date.today().strftime("%d-%m-%Y")
        empty_df.to_excel(file_path, index=False, sheet_name=today_str)

    today_str = datetime.date.today().strftime("%d-%m-%Y")
    sheet_name=f"{today_str}"

    try:
        repairs_excel_df = pd.read_excel(file_path, sheet_name=sheet_name)
    except ValueError:
        empty_df = pd.DataFrame(columns=["No.", "Area", "Vehicle ID", "Date", "Description", "Cost (ugx)"])
        with pd.ExcelWriter(file_path, engine="openpyxl", mode="a", if_sheet_exists="overlay") as writer:
            empty_df.to_excel(writer, index=False, sheet_name=sheet_name)
        repairs_excel_df = empty_df.copy()

    edited_df = st.data_editor(
        repairs_excel_df,
        num_rows="dynamic",
        width="stretch",
        key="data_editor"
    )

    col1, col2, col3 = st.columns([1, 1, 1])

    with col1:
        if st.button("Save Changes"):
            monthly_repairs_df = edited_df.copy()
            repairs_file_path = os.path.join("static", "repairs_excel.xlsx")
            today_str = datetime.date.today().strftime("%d-%m-%Y")
            sheet_name = f"{today_str}"


            try:
                with pd.ExcelWriter(repairs_file_path, engine="openpyxl", mode="a", if_sheet_exists="replace") as writer:
                    monthly_repairs_df.to_excel(writer, index=False, sheet_name=sheet_name)
            except FileNotFoundError:
                with pd.ExcelWriter(repairs_file_path, engine="openpyxl") as writer:
                    monthly_repairs_df.to_excel(writer, index=False, sheet_name=sheet_name)

            work_book = load_workbook(repairs_file_path)
            ws = work_book[sheet_name]

            thin_border = Border(
                left=Side(style='thin'),
                right=Side(style='thin'),
                top=Side(style='thin'),
                bottom=Side(style='thin')
            )
            col_max_width = {}

            for row in ws.iter_rows():
                for cell in row:
                    cell.alignment = Alignment(wrap_text=True, vertical="top")
                    cell.border = thin_border
                    val = str(cell.value) if cell.value is not None else ""
                    col_letter = cell.column_letter
                    col_max_width[col_letter] = max(col_max_width.get(col_letter, 0), len(val))

            # Set column widths based on max content length
            for col_letter, max_len in col_max_width.items():
                ws.column_dimensions[col_letter].width = max(15, min(max_len + 3, 60))  # reasonable min/max
            work_book.save(repairs_file_path)
            st.success("Changes saved successfully!")
        
    with col2:
        if st.button("generate_request"):
            # Computing total cost
            file_path = os.path.join("static", "repairs_excel.xlsx")
            today_str = datetime.date.today().strftime("%d-%m-%Y")
            sheet_name=f"{today_str}"
            monthly_repairs_df = pd.read_excel(file_path, sheet_name=sheet_name)
            monthly_repairs_df["Cost (ugx)"] = monthly_repairs_df["Cost (ugx)"].replace({',': ''}, regex=True).astype(int)

            total_rows = monthly_repairs_df[monthly_repairs_df['Description'] == 'Total Cost (ugx)']

            total_vehicles = int(monthly_repairs_df["Vehicle ID"].nunique())

            if not total_rows.empty:
                total_cost = total_rows["Cost (ugx)"].sum()
            else:
                total_cost = 0

            doc = Document()
            doc.styles['Normal'].font.name = 'Times New Roman'
            doc.styles['Normal'].font.size = Pt(12)

            # Heading (Optional Logo/Header Section)
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run1 = p.add_run("MID-WESTERN UMBRELLA OF WATER AND SANITATION\n")
            run1.bold = True
            run1.font.size = Pt(14)

            run2 = p.add_run("MEMO\n")
            run2.bold = True
            run2.font.size = Pt(14)

            run3 = p.add_run("=" * 60)
            run3.font.size = Pt(12)
            run3.bold = True

            # Recipient block
            to_section = doc.add_paragraph()
            to_section.alignment = WD_ALIGN_PARAGRAPH.LEFT
            to_section.add_run("To:\t\tManager UWS-MW\n")
            to_section.add_run("Thru:\t\tAdministrator\n")
            to_section.add_run("From:\t\tEngineer UWS-MW\n")

            today_str = datetime.date.today().strftime("%d/%B/%Y")
            to_section.add_run(f"Date:\t\t{today_str}")
            to_section.add_run("\n")


            # Subject line
            subj = doc.add_paragraph(f"SUBJECT:\tREQUEST FOR UGSHS {total_cost:,} /= "
                                    f"({num2words(int(total_cost), lang='en').upper()} UGANDA SHILLINGS ONLY) "
                                    f"TO BE PAID TO A&B MOTORCYCLE GARAGE "
                                    f"FOR MOTORCYCLE MAINTENANCE SERVICES PROVIDED FOR NO. {total_vehicles} ({num2words(total_vehicles, lang='en').upper()}) MOTORCYCLES.")
            subj.runs[0].bold = True
            # Underline the subject line
            for run in subj.runs:
                run.underline = True

            # Body text
            body = doc.add_paragraph()
            body.add_run(
                "This is to request the release of the above-mentioned funds, intended for payment to "
                "A&B Motorcycle Garage (Brian Asiimwe) for conducting follow-up repairs and servicing "
                f"on No. {total_vehicles} motorcycles allocated across various areas under the Mid-Western Umbrella for the previous month.\n\n"
                "Management resolved to engage Brian Asiimwe, our currently assigned mechanic, to visit all No. 16 operational areas every month, "
                "to enhance the mechanical condition and overall welfare of the motorcycles used by scheme staff.\n\n"
                "Details of each motorcycle, including the Vehicle ID numbers, total charges, and dates of repair/servicing, have been summarized "
                "in the table below:"
            )

            # Table with details
            table_with_vehicle_details = monthly_repairs_df[monthly_repairs_df["Vehicle ID"].notna() & (monthly_repairs_df["Vehicle ID"] != "")]
            table_with_total_cost = monthly_repairs_df[monthly_repairs_df["Description"] == "Total Cost (ugx)"]

            summary_list = []
            # Prepare summary data for the table
            for i in range(len(table_with_vehicle_details)):
                summary_list.append({
                    "No.": int(table_with_vehicle_details.iloc[i]["No."]),
                    "Area": table_with_vehicle_details.iloc[i]["Area"],
                    "Vehicle ID": table_with_vehicle_details.iloc[i]["Vehicle ID"],
                    "Date": pd.to_datetime(table_with_vehicle_details.iloc[i]["Date"]).strftime("%d, %b, %Y"),
                    "Total Cost (ugx)": f"{int(table_with_total_cost.iloc[i]['Cost (ugx)']):,}" if not table_with_total_cost.empty else ""
                })

            total_amount = table_with_total_cost["Cost (ugx)"].sum()
            summary_list.append({
                "No.": "",
                "Area": "",
                "Vehicle ID": "",
                "Date": "Total Amount (ugx)",
                "Total Cost (ugx)": f"{total_amount:,}" if total_amount else ""
            })

            if summary_list:
                table = doc.add_table(rows=1, cols=len(summary_list[0]))
                table.style = 'Table Grid'
                hdr_cells = table.rows[0].cells
                for idx, key in enumerate(summary_list[0].keys()):
                    hdr_cells[idx].text = key

                for row in summary_list:
                    max_length = max(len(str(v)) for v in row.values())
                    if max_length > 300:
                        for cell in table.rows[-1].cells:
                            cell.width = Pt(300)
                    else:
                        for cell in table.rows[-1].cells:
                            cell.width = Pt(max_length + 20)
                    row_cells = table.add_row().cells
                    for idx, key in enumerate(row.keys()):
                        row_cells[idx].text = str(row[key])

            body2 = doc.add_paragraph()
            body2.add_run(
                "\nThe individual costs for each motorcycle repair and service are further broken down as follows; "
            )

            detailed_list = []
            for i in range(len(monthly_repairs_df)):
                detailed_list.append({
                    "No.": int(monthly_repairs_df.iloc[i]["No."]) if pd.notna(monthly_repairs_df.iloc[i]["No."]) else "",
                    "Area": monthly_repairs_df.iloc[i]["Area"] if pd.notna(monthly_repairs_df.iloc[i]["Area"]) else "",
                    "Vehicle ID": monthly_repairs_df.iloc[i]["Vehicle ID"] if pd.notna(monthly_repairs_df.iloc[i]["Vehicle ID"]) else "",
                    "Date": pd.to_datetime(monthly_repairs_df.iloc[i]["Date"]).strftime("%d, %b, %Y") if pd.notna(monthly_repairs_df.iloc[i]["Date"]) and monthly_repairs_df.iloc[i]["Date"] != "" else "",
                    "Description": monthly_repairs_df.iloc[i]["Description"] if pd.notna(monthly_repairs_df.iloc[i]["Description"]) else "",
                    "Cost (ugx)": f"{int(monthly_repairs_df.iloc[i]['Cost (ugx)']):,}" if pd.notna(monthly_repairs_df.iloc[i]['Cost (ugx)']) and monthly_repairs_df.iloc[i]['Cost (ugx)'] != "" else ""
                })

            if detailed_list:
                table = doc.add_table(rows=1, cols=len(detailed_list[0]))
                table.style = 'Table Grid'
                hdr_cells = table.rows[0].cells
                for idx, key in enumerate(detailed_list[0].keys()):
                    hdr_cells[idx].text = key

                for row in detailed_list:
                    max_length = max(len(str(v)) for v in row.values())
                    if max_length > 300:
                        for cell in table.rows[-1].cells:
                            cell.width = Pt(300)
                    else:
                        for cell in table.rows[-1].cells:
                            cell.width = Pt(max_length + 20)
                    row_cells = table.add_row().cells
                    for idx, key in enumerate(row.keys()):
                        row_cells[idx].text = str(row[key])

            doc.add_paragraph("\nMARVIN LUYOMBYA")
            doc.add_paragraph("\nENGINEER, MWUWS")

            # Save to stream
            doc_io = BytesIO()
            doc.save(doc_io)
            doc_io.seek(0)

            # Save the generated Word document to a folder
            output_folder = os.path.join("static", "generated_requests")
            os.makedirs(output_folder, exist_ok=True)
            today_str = datetime.date.today().strftime("%d-%b-%Y")
            output_path = os.path.join(output_folder, f"repair_request_{today_str}.docx")
            with open(output_path, "wb") as f:
                f.write(doc_io.getbuffer())
            st.success("Repair request document generated!")


    with col3:
        if st.button("update_vehicle_records"):

            repairs_file_path = os.path.join("static", "repairs_excel.xlsx")
            today_str = datetime.date.today().strftime("%d-%m-%Y")
            sheet_name = f"{today_str}"

            monthly_repairs_df = pd.read_excel(os.path.join("static", "repairs_excel.xlsx"), sheet_name=sheet_name)
            monthly_repairs_df = monthly_repairs_df[monthly_repairs_df["Description"] != "Total Cost (ugx)"]
            monthly_repairs_df["No."] = monthly_repairs_df["No."].fillna(method="ffill")
            monthly_repairs_df["Area"] = monthly_repairs_df["Area"].fillna(method="ffill")
            monthly_repairs_df["Vehicle ID"] = monthly_repairs_df["Vehicle ID"].fillna(method="ffill")
            monthly_repairs_df["Date"] = monthly_repairs_df["Date"].fillna(method="ffill")

            grouped_data = monthly_repairs_df.groupby(["No."])

            repair_data = []
            for group_key, group in grouped_data:
                pass
                descriptions = ", ".join(group["Description"].dropna().astype(str).tolist())
                total_cost = group["Cost (ugx)"].replace({',': ''}, regex=True).astype(int).sum()
                repair_data.append({
                    "Area": group["Area"].iloc[0],
                    "Vehicle ID": group["Vehicle ID"].iloc[0],
                    "Date": pd.to_datetime(group["Date"].iloc[0]).strftime("%d-%b-%Y"),
                    "Descriptions": descriptions,
                    "Total Cost (ugx)": f"{total_cost:,}"
                })
            repair_df = pd.DataFrame(repair_data)
            st.session_state['repair_df'] = repair_df

            history_file = os.path.join("static", "repair_history.xlsx")

            if not os.path.exists(history_file):
                columns = ["Area", "Vehicle ID", "Date", "Descriptions", "Total Cost (ugx)"]
                empty_df = pd.DataFrame(columns=columns)
                with pd.ExcelWriter(history_file, engine="openpyxl") as writer:
                    empty_df.to_excel(writer, index=False)

            with pd.ExcelWriter(history_file, engine="openpyxl", mode="a", if_sheet_exists="overlay") as writer:
                existing_df = pd.read_excel(history_file)
                combined_df = pd.concat([existing_df, repair_df], ignore_index=True)
                combined_df.drop_duplicates(subset=["Area", "Vehicle ID", "Date", "Descriptions", "Total Cost (ugx)"], inplace=True)
                combined_df.to_excel(writer, index=False)

            work_book = load_workbook(history_file)
            ws = work_book.active

            thin_border = Border(
                left=Side(style='thin'),
                right=Side(style='thin'),
                top=Side(style='thin'),
                bottom=Side(style='thin')
            )
            col_max_width = {}

            for row in ws.iter_rows():
                for cell in row:
                    cell.alignment = Alignment(wrap_text=True, vertical="top")
                    cell.border = thin_border
                    val = str(cell.value) if cell.value is not None else ""
                    col_letter = cell.column_letter
                    col_max_width[col_letter] = max(col_max_width.get(col_letter, 0), len(val))

            # Set column widths based on max content length
            for col_letter, max_len in col_max_width.items():
                ws.column_dimensions[col_letter].width = max(15, min(max_len + 3, 60))  # reasonable min/max

            work_book.save(history_file)
            st.success("Vehicle repair history updated!")

